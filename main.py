"""
Document Extraction Service
============================
OCR engine: EasyOCR + PyMuPDF + OpenCV  (Tesseract/Poppler yoxdur)

POST /extract-documents   — multipart upload (sender + receiver faylı)
POST /extract-by-url      — JSON body { sender_url, receiver_url }

Response: { sender: { filename, text, lines[] }, receiver: { ... } }
"""

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import fitz                          # PyMuPDF — Poppler lazım deyil
import easyocr
import numpy as np
import cv2
from PIL import Image
from docx import Document
import httpx
import io
import os
import uvicorn

# ── EasyOCR reader (startup-da yüklənir, sonra paylaşılır) ───────────────────
# Lazım olan dilləri əlavə et, məs: ['en', 'az'] — lakin 'az' yoxdur, 'en' kifayətdir
reader = easyocr.Reader(['en'], gpu=False)

app = FastAPI(title="Document Extraction Service", version="2.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["POST", "GET"],
    allow_headers=["*"],
)


# ── Core OCR helpers (sənin işlək kodun) ─────────────────────────────────────

def _preprocess(img: np.ndarray) -> np.ndarray:
    """BGR/GRAY image → contrast artırılmış grayscale."""
    if len(img.shape) == 3:
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    else:
        gray = img
    gray = cv2.equalizeHist(gray)
    return gray


def _ocr(img: np.ndarray) -> list[str]:
    """EasyOCR ilə image-dən text sətirləri qaytar."""
    result = reader.readtext(img)
    return [r[1] for r in result]


def _bytes_to_bgr(data: bytes) -> np.ndarray:
    """Raw bytes → OpenCV BGR array (PIL vasitəsilə)."""
    pil = Image.open(io.BytesIO(data)).convert("RGB")
    return cv2.cvtColor(np.array(pil), cv2.COLOR_RGB2BGR)


# ── Extractors ────────────────────────────────────────────────────────────────

def extract_from_pdf(data: bytes) -> tuple[list[str], str]:
    """PDF hər səhifəsini 300 DPI-da render edib EasyOCR ilə oxu."""
    doc = fitz.open(stream=data, filetype="pdf")
    all_lines: list[str] = []
    for page in doc:
        pix = page.get_pixmap(dpi=300)
        # PyMuPDF samples → numpy array
        img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
            pix.height, pix.width, pix.n
        )
        if pix.n == 4:                         # RGBA → BGR
            img = cv2.cvtColor(img, cv2.COLOR_RGBA2BGR)
        img = _preprocess(img)
        all_lines.extend(_ocr(img))
    lines = [t.strip() for t in all_lines if t.strip()]
    return lines, "\n".join(lines)


def extract_from_image(data: bytes) -> tuple[list[str], str]:
    """JPG / PNG → preprocess → EasyOCR."""
    img = _bytes_to_bgr(data)
    img = _preprocess(img)
    lines = [t.strip() for t in _ocr(img) if t.strip()]
    return lines, "\n".join(lines)


def extract_from_docx(data: bytes) -> tuple[list[str], str]:
    """DOCX → python-docx ilə plain text (OCR lazım deyil)."""
    doc = Document(io.BytesIO(data))
    lines: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return lines, "\n".join(lines)


def extract_text(filename: str, data: bytes) -> dict:
    """Uzantıya görə extractor seç → { lines, text } qaytar."""
    ext = os.path.splitext(filename.lower())[1]
    if ext == ".pdf":
        lines, text = extract_from_pdf(data)
    elif ext in (".jpg", ".jpeg", ".png"):
        lines, text = extract_from_image(data)
    elif ext in (".docx",):
        lines, text = extract_from_docx(data)
    elif ext == ".doc":
        return {
            "lines": [],
            "text": "[Köhnə .doc formatı dəstəklənmir — DOCX olaraq yükləyin]",
        }
    else:
        raise HTTPException(
            status_code=415,
            detail=f"Dəstəklənməyən format: {ext}. PDF, JPG, PNG, DOCX göndərin.",
        )
    return {"lines": lines, "text": text}


# ── Endpoints ──────────────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    return {"status": "ok", "service": "document-extraction", "ocr": "easyocr+pymupdf"}


# ── URL-based endpoint (test üçün) ────────────────────────────────────────────

class ExtractByUrlRequest(BaseModel):
    sender_url: str
    receiver_url: str


@app.post("/extract-by-url")
async def extract_by_url(body: ExtractByUrlRequest):
    """
    Test üçün: JSON body ilə iki URL göndər.

        { "sender_url": "https://...", "receiver_url": "https://..." }

    Response:
        {
          "sender":   { "url", "content_type", "size_bytes", "lines": [...], "text": "..." },
          "receiver": { ... }
        }
    """
    results = {}

    async with httpx.AsyncClient(timeout=60.0, follow_redirects=True) as client:
        for role, url in [("sender", body.sender_url), ("receiver", body.receiver_url)]:
            try:
                resp = await client.get(url)
                resp.raise_for_status()
            except httpx.HTTPError as e:
                raise HTTPException(400, detail=f"{role} URL yüklənmədi: {e}")

            data = resp.content
            if not data:
                raise HTTPException(400, detail=f"{role} URL boş cavab qaytardı.")

            content_type = resp.headers.get("content-type", "")
            ext = _ext_from_content_type(content_type) or _ext_from_url(url)

            try:
                extracted = extract_text(f"file{ext}", data)
            except HTTPException:
                raise
            except Exception as e:
                raise HTTPException(500, detail=f"{role} oxunarkən xəta: {e}")

            results[role] = {
                "url":          url,
                "content_type": content_type,
                "size_bytes":   len(data),
                "lines":        extracted["lines"],
                "text":         extracted["text"],
            }

    return results


# ── Multipart upload endpoint ─────────────────────────────────────────────────

@app.post("/extract-documents")
async def extract_documents(
    sender_document:   UploadFile = File(...),
    receiver_document: UploadFile = File(...),
):
    """
    İki fayl yüklə (multipart/form-data).

    Response:
        {
          "sender":   { "filename", "content_type", "size_bytes", "lines": [...], "text": "..." },
          "receiver": { ... }
        }
    """
    results = {}

    for role, upload in [("sender", sender_document), ("receiver", receiver_document)]:
        data = await upload.read()
        if not data:
            raise HTTPException(400, detail=f"{role} sənədi boşdur.")

        try:
            extracted = extract_text(upload.filename or f"{role}.pdf", data)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(500, detail=f"{role} oxunarkən xəta: {e}")

        results[role] = {
            "filename":     upload.filename,
            "content_type": upload.content_type,
            "size_bytes":   len(data),
            "lines":        extracted["lines"],
            "text":         extracted["text"],
        }

    return results


# ── Helpers ───────────────────────────────────────────────────────────────────

def _ext_from_content_type(ct: str) -> str:
    ct = ct.lower().split(";")[0].strip()
    return {
        "application/pdf":   ".pdf",
        "image/jpeg":        ".jpg",
        "image/png":         ".png",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "application/msword": ".doc",
    }.get(ct, "")


def _ext_from_url(url: str) -> str:
    path = url.split("?")[0]
    _, ext = os.path.splitext(path)
    return ext.lower() or ".pdf"


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)